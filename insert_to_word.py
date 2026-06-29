#!/usr/bin/env python3
"""
Script untuk menggabungkan gambar/PDF/Word dari folder ke dalam dokumen Word.
- Auto-deteksi tipe file (image/pdf/docx)
- Convert PDF/Word → image (screenshot halaman pertama)
- Ukuran gambar: 70% dari aslinya
- Header 1: "Lampiran"
- Scan recursive subfolder
"""
import os
import sys
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: pdf2image tidak terinstall. PDF/Word tidak akan diproses.")
    print("Install dengan: pip install pdf2image")


def create_lampiran_header(doc, num, name):
    """Buat header level 1 untuk lampiran"""
    header = doc.add_heading(f"Lampiran {num}: {name}", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return header


def embed_image_from_file(doc, image_path, size_percent=0.5):
    """Embed image dengan ukuran 70% dari aslinya"""
    try:
        img = Image.open(image_path)
        width, height = img.size
        new_width = int(width * size_percent)
        new_height = int(height * size_percent)
        img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img_resized.save(tmp.name, 'PNG')
            temp_path = tmp.name

        doc.add_paragraph()
        width_inches = new_width / 96.0
        doc.add_picture(temp_path, width=Inches(min(width_inches, 6.5)))
        doc.add_paragraph()

        os.remove(temp_path)
        return True
    except Exception as e:
        print(f"  Error embedding image {image_path}: {e}")
        return False


def convert_pdf_to_image(pdf_path, output_image_path):
    """Convert PDF ke image (halaman pertama)"""
    if not PDF_SUPPORT:
        return False

    try:
        pages = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=1)
        if pages:
            pages[0].save(output_image_path, 'PNG')
            return True
        return False
    except Exception as e:
        print(f"  Error converting PDF to image: {e}")
        return False


def convert_word_to_pdf(word_path, pdf_path):
    """Extract text from Word dan convert ke simple PDF (fallback tanpa LibreOffice)"""
    try:
        doc = Document(word_path)

        # Buat PDF document baru
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        y = height - 40

        # Tulis text dari paragraf
        for para in doc.paragraphs:
            if para.text.strip():
                c.drawString(40, y, para.text[:80])  # Limit text length
                y -= 20
                if y < 40:
                    c.showPage()
                    y = height - 40

        c.save()
        return True
    except ImportError:
        print(f"  reportlab tidak terinstall. Install dengan: pip install reportlab")
        return False
    except Exception as e:
        print(f"  Error converting Word to PDF: {e}")
        return False


def embed_pdf(doc, pdf_path, size_percent=0.7):
    """Embed PDF sebagai image (screenshot halaman pertama)"""
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        temp_img = tmp.name

    if convert_pdf_to_image(pdf_path, temp_img):
        success = embed_image_from_file(doc, temp_img, size_percent)
        if os.path.exists(temp_img):
            os.remove(temp_img)
        return success

    if os.path.exists(temp_img):
        os.remove(temp_img)
    return False


def embed_word(doc, word_path, size_percent=0.7):
    """Embed Word sebagai image (convert → PDF → image)"""
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
        temp_pdf = tmp_pdf.name

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
        temp_img = tmp_img.name

    try:
        # Convert Word → PDF
        if not convert_word_to_pdf(word_path, temp_pdf):
            return False

        # Convert PDF → image
        if convert_pdf_to_image(temp_pdf, temp_img):
            success = embed_image_from_file(doc, temp_img, size_percent)
            return success

        return False
    finally:
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)
        if os.path.exists(temp_img):
            os.remove(temp_img)


def scan_files_recursive(folder_path):
    """Scan recursive semua file image/pdf/docx"""
    image_exts = {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff',
                  'JPG', 'JPEG', 'PNG', 'GIF', 'BMP', 'TIFF'}
    pdf_exts = {'pdf', 'PDF'}
    word_exts = {'docx', 'DOCX', 'doc', 'DOC'}

    all_files = []

    for root, dirs, files in os.walk(folder_path):
        for f in files:
            ext = f.rsplit('.', 1)[-1] if '.' in f else ''
            if ext in image_exts or ext in pdf_exts or ext in word_exts:
                all_files.append(os.path.join(root, f))

    return sorted(all_files)


def process_folder(folder_path, output_file='dokumentasi_penelitian.docx'):
    """Proses folder dan generate dokumen Word"""

    if not os.path.exists(folder_path):
        print(f"Error: Folder tidak ditemukan: {folder_path}")
        return

    doc = Document()
    doc.add_heading('Dokumentasi Penelitian', level=0)
    doc.add_paragraph()

    all_files = scan_files_recursive(folder_path)

    if not all_files:
        print(f"Tidak ada file gambar, PDF, atau Word ditemukan di: {folder_path}")
        return

    print(f"Ditemukan {len(all_files)} file untuk diproses...\n")

    lampiran_num = 1
    success_count = 0

    for file_path in all_files:
        rel_path = os.path.relpath(file_path, folder_path)
        filename = os.path.basename(file_path)
        ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''

        print(f"Memproses [{lampiran_num}]: {rel_path}")

        create_lampiran_header(doc, lampiran_num, filename)

        if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'):
            success = embed_image_from_file(doc, file_path)
        elif ext == 'pdf':
            success = embed_pdf(doc, file_path)
        elif ext in ('docx', 'doc'):
            success = embed_word(doc, file_path)
        else:
            print(f"  Skip: tipe file tidak dikenal")
            success = False

        if success:
            success_count += 1

        lampiran_num += 1

    doc.save(output_file)
    print(f"\n{'='*50}")
    print(f"Dokumen Word berhasil dibuat: {output_file}")
    print(f"Total file diproses: {success_count}/{len(all_files)}")
    print(f"{'='*50}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python insert_to_word.py <folder_path> [output_file.docx]")
        print("\nContoh:")
        print("  python insert_to_word.py '/path/to/dokumentasi penelitian'")
        print("  python insert_to_word.py '/path/to/dokumentasi penelitian' output.docx")
        sys.exit(1)

    folder_path = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'dokumentasi_penelitian.docx'

    process_folder(folder_path, output_file)
